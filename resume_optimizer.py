import os
import io
import re
import threading
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
import customtkinter as ctk
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
import openai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set up OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Set CustomTkinter appearance
ctk.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
ctk.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

class ResumeOptimizerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Resume Optimizer Pro")
        self.root.geometry("1000x750")
        self.root.minsize(900, 700)
        
        # Theme colors
        self.colors = {
            "primary": "#1E88E5",
            "secondary": "#26A69A",
            "accent": "#FF8F00",
            "background": "#F5F7FA",
            "card": "#FFFFFF",
            "text_primary": "#212121",
            "text_secondary": "#757575",
        }
        
        # Variables
        self.resume_path = ""
        self.resume_format = ""
        self.original_resume_content = None
        self.job_description = ""
        self.optimized_resume_text = ""
        self.analysis_report = ""
        
        # Create main frame with improved appearance
        self.main_frame = ctk.CTkFrame(self.root, fg_color=self.colors["background"], corner_radius=0)
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create scrollable container
        self.scrollable_frame = ctk.CTkScrollableFrame(self.main_frame, fg_color=self.colors["background"])
        self.scrollable_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # App title and header
        self.header_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.colors["card"], corner_radius=10)
        self.header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.title_label = ctk.CTkLabel(
            self.header_frame, 
            text="Resume Optimizer Pro", 
            font=ctk.CTkFont(size=28, weight="bold"),
            text_color=self.colors["primary"]
        )
        self.title_label.pack(pady=(20, 5))
        
        self.desc_label = ctk.CTkLabel(
            self.header_frame, 
            text="Tailor your resume to job descriptions and boost your ATS match score",
            font=ctk.CTkFont(size=16),
            text_color=self.colors["text_secondary"]
        )
        self.desc_label.pack(pady=(0, 20))
        
        # Input card with shadow effect
        self.input_card = ctk.CTkFrame(self.scrollable_frame, fg_color=self.colors["card"], corner_radius=10)
        self.input_card.pack(fill=tk.X, padx=10, pady=10)
        
        # Section title
        self.input_title = ctk.CTkLabel(
            self.input_card, 
            text="Job & Resume Details", 
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=self.colors["text_primary"]
        )
        self.input_title.pack(anchor="w", padx=25, pady=(20, 15))
        
        # Input fields container
        self.input_fields_frame = ctk.CTkFrame(self.input_card, fg_color="transparent")
        self.input_fields_frame.pack(fill=tk.X, padx=25, pady=(0, 20))
        
        # Job URL
        self.url_label = ctk.CTkLabel(
            self.input_fields_frame, 
            text="Job Posting URL:",
            font=ctk.CTkFont(size=14),
            text_color=self.colors["text_primary"]
        )
        self.url_label.grid(row=0, column=0, sticky="w", padx=(0, 15), pady=(0, 15))
        
        self.url_entry = ctk.CTkEntry(
            self.input_fields_frame, 
            width=400,
            height=40,
            placeholder_text="Paste job posting URL here"
        )
        self.url_entry.grid(row=0, column=1, sticky="ew", padx=0, pady=(0, 15))
        
        # Alternative text input for job description
        self.alt_input_label = ctk.CTkLabel(
            self.input_fields_frame, 
            text="Or paste job description:",
            font=ctk.CTkFont(size=14),
            text_color=self.colors["text_primary"]
        )
        self.alt_input_label.grid(row=1, column=0, sticky="nw", padx=(0, 15), pady=(0, 15))
        
        self.job_desc_text = ctk.CTkTextbox(
            self.input_fields_frame,
            width=400,
            height=100
        )
        self.job_desc_text.grid(row=1, column=1, sticky="ew", padx=0, pady=(0, 15))
        
        # Resume file selection
        self.file_label = ctk.CTkLabel(
            self.input_fields_frame, 
            text="Resume File (PDF/DOCX):",
            font=ctk.CTkFont(size=14),
            text_color=self.colors["text_primary"]
        )
        self.file_label.grid(row=2, column=0, sticky="w", padx=(0, 15), pady=0)
        
        self.file_frame = ctk.CTkFrame(self.input_fields_frame, fg_color="transparent")
        self.file_frame.grid(row=2, column=1, sticky="ew", padx=0, pady=0)
        
        self.file_entry = ctk.CTkEntry(
            self.file_frame, 
            width=300,
            height=40,
            placeholder_text="Select your resume file"
        )
        self.file_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 10))
        
        self.browse_button = ctk.CTkButton(
            self.file_frame, 
            text="Browse",
            height=40,
            fg_color=self.colors["secondary"],
            hover_color=self.colors["secondary"] + "C0",
            command=self.browse_file
        )
        self.browse_button.pack(side=tk.RIGHT)
        
        # Output format selection
        self.format_label = ctk.CTkLabel(
            self.input_fields_frame, 
            text="Output Format:",
            font=ctk.CTkFont(size=14),
            text_color=self.colors["text_primary"]
        )
        self.format_label.grid(row=3, column=0, sticky="w", padx=(0, 15), pady=(15, 0))
        
        self.format_frame = ctk.CTkFrame(self.input_fields_frame, fg_color="transparent")
        self.format_frame.grid(row=3, column=1, sticky="ew", padx=0, pady=(15, 0))
        
        self.format_var = tk.StringVar(value="same_as_input")
        formats = [
            ("Same as input", "same_as_input"),
            ("Plain text (.txt)", "txt"),
            ("Word (.docx)", "docx"),
            ("PDF (.pdf)", "pdf")
        ]
        
        for i, (text, value) in enumerate(formats):
            radio = ctk.CTkRadioButton(
                self.format_frame,
                text=text,
                value=value,
                variable=self.format_var,
                text_color=self.colors["text_primary"],
                fg_color=self.colors["primary"]
            )
            radio.pack(side=tk.LEFT, padx=(0 if i == 0 else 20, 0))
        
        # Configure grid column weights
        self.input_fields_frame.columnconfigure(1, weight=1)
        
        # Process button with modern design
        self.button_frame = ctk.CTkFrame(self.scrollable_frame, fg_color="transparent")
        self.button_frame.pack(fill=tk.X, padx=10, pady=(20, 10))
        
        self.process_button = ctk.CTkButton(
            self.button_frame, 
            text="Optimize Resume", 
            font=ctk.CTkFont(size=16, weight="bold"),
            height=50,
            corner_radius=25,
            fg_color=self.colors["primary"],
            hover_color=self.colors["primary"] + "C0",
            command=self.process_resume
        )
        self.process_button.pack(pady=10, padx=100)
        
        # Progress indicator with animation
        self.progress_frame = ctk.CTkFrame(self.scrollable_frame, fg_color="transparent")
        self.progress_frame.pack(fill=tk.X, padx=10)
        
        self.progress_bar = ctk.CTkProgressBar(self.progress_frame)
        self.progress_bar.pack(fill=tk.X, padx=100, pady=(0, 5))
        self.progress_bar.set(0)
        
        self.progress_label = ctk.CTkLabel(
            self.progress_frame, 
            text="",
            font=ctk.CTkFont(size=14),
            text_color=self.colors["text_secondary"]
        )
        self.progress_label.pack(pady=(0, 10))
        
        # Hide progress indicator initially
        self.progress_frame.pack_forget()
        
        # Results section
        self.results_frame = ctk.CTkFrame(self.scrollable_frame, fg_color="transparent")
        self.results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Notebook-style tabs for results with improved styling
        self.tab_view = ctk.CTkTabview(
            self.results_frame,
            fg_color=self.colors["card"],
            segmented_button_fg_color=self.colors["background"],
            segmented_button_selected_color=self.colors["primary"],
            segmented_button_selected_hover_color=self.colors["primary"] + "C0",
            segmented_button_unselected_hover_color=self.colors["background"] + "C0"
        )
        self.tab_view.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # Add tabs
        self.tab_view.add("Analysis Report")
        self.tab_view.add("Optimized Resume")
        
        # Analysis report text area with improved styling
        self.analysis_text = scrolledtext.ScrolledText(
            self.tab_view.tab("Analysis Report"),
            wrap=tk.WORD,
            font=("Arial", 12),
            bg=self.colors["card"],
            fg=self.colors["text_primary"],
            insertbackground=self.colors["primary"],
            padx=15,
            pady=15
        )
        self.analysis_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Optimized resume text area
        self.resume_text = scrolledtext.ScrolledText(
            self.tab_view.tab("Optimized Resume"),
            wrap=tk.WORD,
            font=("Arial", 12),
            bg=self.colors["card"],
            fg=self.colors["text_primary"],
            insertbackground=self.colors["primary"],
            padx=15,
            pady=15
        )
        self.resume_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Save buttons frame with modern styling
        self.save_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.colors["card"], corner_radius=10)
        self.save_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.save_buttons_title = ctk.CTkLabel(
            self.save_frame, 
            text="Save Results", 
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=self.colors["text_primary"]
        )
        self.save_buttons_title.pack(anchor="w", padx=25, pady=(20, 15))
        
        self.save_buttons_container = ctk.CTkFrame(self.save_frame, fg_color="transparent")
        self.save_buttons_container.pack(fill=tk.X, padx=25, pady=(0, 20))
        
        self.save_analysis_button = ctk.CTkButton(
            self.save_buttons_container,
            text="Save Analysis Report",
            height=40,
            fg_color=self.colors["secondary"],
            hover_color=self.colors["secondary"] + "C0",
            command=lambda: self.save_text_to_file(self.analysis_text.get("1.0", tk.END), "analysis_report.txt")
        )
        self.save_analysis_button.pack(side=tk.LEFT, padx=(0, 20), pady=10)
        
        self.save_resume_button = ctk.CTkButton(
            self.save_buttons_container,
            text="Save as Plain Text",
            height=40,
            fg_color=self.colors["secondary"],
            hover_color=self.colors["secondary"] + "C0",
            command=lambda: self.save_text_to_file(self.resume_text.get("1.0", tk.END), "optimized_resume.txt")
        )
        self.save_resume_button.pack(side=tk.LEFT, padx=(0, 20), pady=10)
        
        self.save_formatted_button = ctk.CTkButton(
            self.save_buttons_container,
            text="Save Formatted Resume",
            height=40,
            fg_color=self.colors["accent"],
            hover_color=self.colors["accent"] + "C0",
            command=self.save_formatted_resume
        )
        self.save_formatted_button.pack(side=tk.LEFT, padx=0, pady=10)
        
        # Initially hide results and save sections
        self.results_frame.pack_forget()
        self.save_frame.pack_forget()
        
        # Check if API key exists
        if not openai.api_key:
            self.show_api_key_dialog()
    
    def show_api_key_dialog(self):
        """Show dialog to enter OpenAI API key if not found in .env"""
        dialog = ctk.CTkInputDialog(
            text="Please enter your OpenAI API key:", 
            title="API Key Required"
        )
        api_key = dialog.get_input()
        if api_key:
            openai.api_key = api_key
            # Create .env file if it doesn't exist
            with open(".env", "w") as f:
                f.write(f"OPENAI_API_KEY={api_key}")
        else:
            messagebox.showerror(
                "API Key Required", 
                "An OpenAI API key is required to use this application. Please restart and enter a valid API key."
            )
            self.root.destroy()
    
    def browse_file(self):
        """Open file dialog to select resume file"""
        filetypes = [
            ("Resume files", "*.pdf;*.docx"),
            ("PDF files", "*.pdf"),
            ("Word files", "*.docx"),
            ("All files", "*.*")
        ]
        filename = filedialog.askopenfilename(filetypes=filetypes)
        if filename:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, filename)
            self.resume_path = filename
            
            # Determine file format
            lower_filename = filename.lower()
            if lower_filename.endswith('.pdf'):
                self.resume_format = 'pdf'
            elif lower_filename.endswith('.docx'):
                self.resume_format = 'docx'
            else:
                self.resume_format = 'unknown'
    
    def save_text_to_file(self, text, default_filename):
        """Save text content to a file"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            initialfile=default_filename,
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        if filename:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(text)
            messagebox.showinfo("Success", f"File saved successfully to {filename}")
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from a PDF file."""
        text = ""
        pdf_content = None
        
        try:
            # Store original content for formatting
            pdf_content = pdf_file.read()
            pdf_file.seek(0)
            
            reader = PyPDF2.PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting text from PDF: {e}")
        
        return text, pdf_content

    def extract_text_from_docx(self, docx_file):
        """Extract text from a DOCX file with formatting information."""
        text = ""
        formatted_text = ""
        docx_content = None
        
        try:
            # Store original content for formatting
            docx_content = docx_file.read()
            docx_file.seek(0)
            
            doc = docx.Document(docx_file)
            
            # Process paragraphs with formatting
            for para in doc.paragraphs:
                # Skip empty paragraphs
                if not para.text.strip():
                    formatted_text += "\n"
                    text += "\n"
                    continue
                
                # Check alignment
                alignment_tag_start = ""
                alignment_tag_end = ""
                
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment_tag_start = "<center>"
                    alignment_tag_end = "</center>"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment_tag_start = "<right>"
                    alignment_tag_end = "</right>"
                
                # Check if the paragraph is a heading
                if para.style.name.startswith('Heading'):
                    formatted_text += f"{alignment_tag_start}**{para.text}**{alignment_tag_end}\n"
                    text += para.text + "\n"
                    continue
                
                # Process runs for formatting within paragraphs
                para_text = ""
                formatted_para = ""
                
                # Check for bullet lists
                if para.style.name.startswith('List'):
                    bullet_prefix = "• "
                else:
                    bullet_prefix = ""
                
                for run in para.runs:
                    run_text = run.text
                    
                    # Check for hyperlinks - this is a simplistic approach
                    hyperlink = False
                    
                    # Apply formatting based on run properties
                    if run.bold and run.italic:
                        formatted_para += f"**_{run_text}_**"
                    elif run.bold:
                        formatted_para += f"**{run_text}**"
                    elif run.italic:
                        formatted_para += f"_{run_text}_"
                    else:
                        formatted_para += run_text
                    
                    para_text += run_text
                
                # Look for metrics (numbers with % or $ signs)
                metrics_pattern = r'\b(\d+%|\$\d+(?:,\d+)*(?:\.\d+)?|\d+\s*%|\d+\s*(?:percent|pct))\b'
                metrics = re.findall(metrics_pattern, para_text, re.IGNORECASE)
                
                # If metrics found, mark them for highlighting
                if metrics:
                    for metric in metrics:
                        # Only replace if not already marked with formatting
                        if metric in formatted_para and f"**[{metric}]**" not in formatted_para:
                            formatted_para = formatted_para.replace(
                                metric, 
                                f"**[{metric}]**"
                            )
                
                # Add bullet if needed
                if bullet_prefix:
                    formatted_para = bullet_prefix + formatted_para
                
                # Add to formatted text
                formatted_text += f"{alignment_tag_start}{formatted_para}{alignment_tag_end}\n"
                text += bullet_prefix + para_text + "\n"
            
            # Get text from tables as well
            for table in doc.tables:
                # Add table marker
                formatted_text += "<table>\n"
                for row in table.rows:
                    row_text = ""
                    formatted_row = ""
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_text += cell_text + "\t"
                        
                        # Check for bold/formatting in cell paragraphs
                        cell_formatted = cell_text
                        
                        # Look for metrics in cells
                        metrics_pattern = r'\b(\d+%|\$\d+(?:,\d+)*(?:\.\d+)?|\d+\s*%|\d+\s*(?:percent|pct))\b'
                        metrics = re.findall(metrics_pattern, cell_text, re.IGNORECASE)
                        
                        # If metrics found, mark them
                        if metrics:
                            for metric in metrics:
                                cell_formatted = cell_formatted.replace(
                                    metric, 
                                    f"**[{metric}]**"
                                )
                        
                        formatted_row += cell_formatted + "\t"
                    
                    text += row_text + "\n"
                    formatted_text += formatted_row + "\n"
                
                formatted_text += "</table>\n"
                text += "\n"
                
        except Exception as e:
            # If advanced parsing fails, fall back to simple extraction
            try:
                docx_file.seek(0)
                doc = docx.Document(docx_file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                formatted_text = text
            except Exception as inner_e:
                messagebox.showerror("Error", f"Error extracting text from DOCX: {inner_e}")
        
        return formatted_text or text, docx_content

    def extract_job_description(self, url):
        """Extract job description from a URL."""
        try:
            response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Try to find common job description containers
            job_description = ""
            
            # Look for job description in various common containers
            job_containers = soup.select('.job-description, .description, .content, [class*="job"], [class*="description"], [id*="job"], [id*="description"]')
            if job_containers:
                # Use the largest container as it's likely to be the main content
                largest_container = max(job_containers, key=lambda x: len(x.get_text()))
                job_description = largest_container.get_text(strip=True, separator='\n')
            else:
                # Fallback to body text if no specific container is found
                job_description = soup.body.get_text(strip=True, separator='\n')
                
            # Clean up the text
            job_description = ' '.join(job_description.split())
                
            return job_description
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting job description: {e}")
            return None

    def optimize_resume(self, resume_text, job_description):
        """Use OpenAI API to optimize resume for the job description."""
        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": """You are an expert resume optimizer who helps candidates match their resumes to job descriptions for better ATS matching scores.
                    
                    When optimizing resumes, you must preserve the original formatting structure:
                    1. Keep section headings exactly as they are (same capitalization, punctuation)
                    2. Maintain heading levels and hierarchy
                    3. Preserve bullet points and list formatting
                    4. Maintain indentation and text alignment patterns
                    5. Preserve any hyperlinks by using [text](url) format
                    6. Keep date formats consistent
                    
                    For formatting instructions:
                    - Use **bold** for headings and important text
                    - Use _italics_ for emphasis where appropriate
                    - Format metrics and key achievements like: **[increased revenue by 25%]**
                    - Preserve center or right alignment with <center> or <right> tags
                    - For left-justified text, no special tags are needed
                    - Maintain any table-like structures by using consistent spacing
                    """},
                    {"role": "user", "content": f"""
                    I need to optimize my resume for a specific job. 
                    
                    Here is my current resume:
                    {resume_text}
                    
                    Here is the job description:
                    {job_description}
                    
                    Please optimize my resume to better match this job description and increase my ATS matching score. 
                    Use the same format as my original resume, but enhance the content to better align with the job requirements.
                    Keep my honest experiences and qualifications, but highlight relevant skills and use appropriate keywords from the job description.
                    
                    Very important formatting requirements:
                    1. Maintain the exact same section structure and headings
                    2. Bold all headings and section titles using **text**
                    3. Make key metrics and achievements bold using **[metric]** format
                    4. Preserve text alignment - use <center> tags for centered text and <right> tags for right-justified text
                    5. Maintain any hyperlinks in the form [text](url)
                    
                    Return only the optimized resume text with thorough formatting indicators.
                    """}
                ],
                temperature=0.2
            )
            return response.choices[0].message.content
        except Exception as e:
            messagebox.showerror("Error", f"Error optimizing resume: {e}")
            return None

    def generate_analysis_report(self, resume_text, job_description):
        """Generate a summary analysis report comparing resume to job description."""
        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert resume analyst who provides objective feedback on resumes."},
                    {"role": "user", "content": f"""
                    I need an analysis of my resume compared to a specific job description.
                    
                    Resume:
                    {resume_text}
                    
                    Job Description:
                    {job_description}
                    
                    Please provide a detailed analysis with the following sections:
                    1. Strengths: What parts of my resume match well with the job description?
                    2. Weaknesses: What important elements from the job description are missing or underrepresented in my resume?
                    3. Areas for Improvement: Specific suggestions to make my resume more competitive for this position.
                    4. Keyword Analysis: Key terms from the job description that should be included in my resume.
                    5. ATS Compatibility Score: Provide an estimated match percentage (0-100%) based on key requirements.
                    
                    Be honest, specific, and actionable in your feedback.
                    """}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            messagebox.showerror("Error", f"Error generating analysis report: {e}")
            return None
    
    def create_pdf_resume(self, filename, content):
        """Create a formatted PDF from the optimized resume text with proper formatting"""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        styles.add(ParagraphStyle(
            name='ResumeHeading',
            parent=styles['Heading1'],
            fontSize=14,
            fontName='Helvetica-Bold',
            spaceAfter=10,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=5,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeBold',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceAfter=5,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeCentered',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_CENTER,
            spaceAfter=5,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeCenteredBold',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=5,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeRight',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_RIGHT,
            spaceAfter=5,
        ))
        
        styles.add(ParagraphStyle(
            name='ResumeRightBold',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            alignment=TA_RIGHT,
            spaceAfter=5,
        ))
        
        # Process content and convert to paragraphs
        story = []
        
        for line in content.split('\n'):
            # Skip empty lines, just add spacing
            if not line.strip():
                story.append(Spacer(1, 10))
                continue
            
            # Check for alignment tags
            alignment_style = 'ResumeNormal'
            processed_line = line
            
            if line.startswith('<center>') and line.endswith('</center>'):
                alignment_style = 'ResumeCentered'
                processed_line = line[8:-9]  # Remove tags
            elif line.startswith('<right>') and line.endswith('</right>'):
                alignment_style = 'ResumeRight'
                processed_line = line[7:-8]  # Remove tags
            
            # Check if line is a heading (bold)
            bold_pattern = r'^\s*\*\*(.*?)\*\*\s*$'
            if re.match(bold_pattern, processed_line):
                heading_text = re.search(bold_pattern, processed_line).group(1)
                story.append(Paragraph(heading_text, styles['ResumeHeading']))
                continue
            
            # Process formatting
            # 1. Convert hyperlinks [text](url)
            hyperlink_pattern = r'\[(.*?)\]\((.*?)\)'
            processed_line = re.sub(
                hyperlink_pattern,
                r'<link href="\2">\1</link>',
                processed_line
            )
            
            # 2. Convert metrics **[metric]** - just make them bold, no highlighting
            metric_pattern = r'\*\*\[(.*?)\]\*\*'
            processed_line = re.sub(
                metric_pattern,
                r'<b>\1</b>',
                processed_line
            )
            
            # 3. Convert bold text **text**
            bold_pattern = r'\*\*(.*?)\*\*'
            processed_line = re.sub(
                bold_pattern,
                r'<b>\1</b>',
                processed_line
            )
            
            # 4. Convert italic text _text_
            italic_pattern = r'_(.*?)_'
            processed_line = re.sub(
                italic_pattern,
                r'<i>\1</i>',
                processed_line
            )
            
            # Determine if the line would require bold style based on content
            needs_bold = '<b>' in processed_line and '</b>' in processed_line
            
            # Select the appropriate style based on alignment and bold content
            if alignment_style == 'ResumeCentered' and needs_bold:
                alignment_style = 'ResumeCenteredBold'
            elif alignment_style == 'ResumeRight' and needs_bold:
                alignment_style = 'ResumeRightBold'
            
            # 5. Handle bullet points
            if processed_line.strip().startswith('•') or processed_line.strip().startswith('-'):
                bullet_text = processed_line.strip()[1:].strip()
                story.append(Paragraph(f"• {bullet_text}", styles[alignment_style]))
            else:
                story.append(Paragraph(processed_line, styles[alignment_style]))
        
        # Build the document
        doc.build(story)
    
    def save_formatted_resume(self):
        """Save the optimized resume with formatting preserved"""
        if not self.optimized_resume_text:
            messagebox.showerror("Error", "No optimized resume to save")
            return
            
        # Determine output format
        output_format = self.format_var.get()
        if output_format == "same_as_input":
            output_format = self.resume_format
        
        # Define potential output file extensions and types
        format_extensions = {
            "txt": ".txt",
            "docx": ".docx",
            "pdf": ".pdf"
        }
        
        format_descriptions = {
            "txt": "Text files",
            "docx": "Word documents",
            "pdf": "PDF files"
        }
        
        # Set default filename and extension based on format
        default_extension = format_extensions.get(output_format, ".txt")
        format_description = format_descriptions.get(output_format, "Text files")
        
        # Ask user for save location
        filename = filedialog.asksaveasfilename(
            defaultextension=default_extension,
            initialfile=f"optimized_resume{default_extension}",
            filetypes=[(format_description, f"*{default_extension}"), ("All files", "*.*")]
        )
        
        if not filename:
            return
        
        try:
            # Create formatted output based on selected format
            if output_format == "txt":
                # For text format, just save as is
                with open(filename, "w", encoding="utf-8") as f:
                    # Remove formatting tags for plain text output
                    clean_text = self.optimized_resume_text
                    # Remove alignment tags
                    clean_text = re.sub(r'<(center|right)>(.*?)</\1>', r'\2', clean_text)
                    # Remove bold markers
                    clean_text = re.sub(r'\*\*\[(.*?)\]\*\*', r'\1', clean_text)
                    clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
                    # Remove italic markers
                    clean_text = re.sub(r'_(.*?)_', r'\1', clean_text)
                    # Convert hyperlinks
                    clean_text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', clean_text)
                    f.write(clean_text)
                    
            elif output_format == "docx":
                # Create a new DOCX with formatting
                doc = docx.Document()
                
                # Process content line by line
                for line in self.optimized_resume_text.split('\n'):
                    if not line.strip():
                        # Skip empty lines
                        doc.add_paragraph()
                        continue
                    
                    # Check for alignment
                    alignment = WD_ALIGN_PARAGRAPH.LEFT  # Default
                    processed_line = line
                    
                    # Handle center alignment
                    if line.startswith('<center>') and line.endswith('</center>'):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER
                        processed_line = line[8:-9]  # Remove tags
                    
                    # Handle right alignment
                    elif line.startswith('<right>') and line.endswith('</right>'):
                        alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        processed_line = line[7:-8]  # Remove tags
                    
                    # Create paragraph with proper alignment
                    para = doc.add_paragraph()
                    para.alignment = alignment
                    
                    # Check if entire line is bold (heading)
                    bold_pattern = r'^\s*\*\*(.*?)\*\*\s*$'
                    if re.match(bold_pattern, processed_line):
                        heading_text = re.search(bold_pattern, processed_line).group(1)
                        run = para.add_run(heading_text)
                        run.bold = True
                        continue
                    
                    # Handle bullet points
                    if processed_line.strip().startswith('•') or processed_line.strip().startswith('-'):
                        para.style = 'List Bullet'
                        processed_line = processed_line.strip()[1:].strip()
                    
                    # We'll now process the remaining text with mixed formatting
                    remaining_text = processed_line
                    
                    while remaining_text:
                        # Look for formatting patterns in this order of precedence
                        patterns = [
                            # Metrics: **[30% increase]** - just make bold, no highlighting
                            (r'\*\*\[(.*?)\]\*\*', lambda m: self._add_highlighted_run(para, m.group(1))),
                            
                            # Hyperlinks: [text](url)
                            (r'\[(.*?)\]\((.*?)\)', lambda m: self._add_hyperlink(doc, para, m.group(1), m.group(2))),
                            
                            # Bold text: **text**
                            (r'\*\*(.*?)\*\*', lambda m: self._add_bold_run(para, m.group(1))),
                            
                            # Italic text: _text_
                            (r'_(.*?)_', lambda m: self._add_italic_run(para, m.group(1)))
                        ]
                        
                        # Find the first matching pattern
                        match = None
                        pattern_index = -1
                        
                        for i, (pattern, _) in enumerate(patterns):
                            m = re.search(pattern, remaining_text)
                            if m and (match is None or m.start() < match.start()):
                                match = m
                                pattern_index = i
                        
                        if match:
                            # Add text before the match
                            if match.start() > 0:
                                para.add_run(remaining_text[:match.start()])
                            
                            # Process the matched pattern
                            _, handler = patterns[pattern_index]
                            handler(match)
                            
                            # Continue with remaining text
                            remaining_text = remaining_text[match.end():]
                        else:
                            # No more patterns, add the rest as plain text
                            if remaining_text:
                                para.add_run(remaining_text)
                            break
                
                # Save the document
                doc.save(filename)
                    
            elif output_format == "pdf":
                # Create PDF with formatting
                self.create_pdf_resume(filename, self.optimized_resume_text)
            
            messagebox.showinfo("Success", f"Formatted resume saved successfully to {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error saving formatted resume: {e}")
    
    def _add_bold_run(self, paragraph, text):
        """Add a bold run to a Word paragraph"""
        run = paragraph.add_run(text)
        run.bold = True
        return text

    def _add_highlighted_run(self, paragraph, text):
        """Add a bold run to a Word paragraph (without highlighting)"""
        run = paragraph.add_run(text)
        run.bold = True
        return text
    
    def _add_italic_run(self, paragraph, text):
        """Add an italic run to a Word paragraph"""
        run = paragraph.add_run(text)
        run.italic = True
        return text

    def _add_hyperlink(self, document, paragraph, text, url):
        """Add a hyperlink to a Word paragraph"""
        try:
            # Create relationship for hyperlink
            rel_id = document.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            
            # Create the hyperlink XML element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), rel_id)
            
            # Create the run element
            run = OxmlElement('w:r')
            
            # Create run properties
            rPr = OxmlElement('w:rPr')
            
            # Set hyperlink style
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Blue
            rPr.append(color)
            
            # Add underline
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            run.append(rPr)
            
            # Add text to the run
            t = OxmlElement('w:t')
            t.text = text
            run.append(t)
            
            # Add the run to the hyperlink
            hyperlink.append(run)
            
            # Add the hyperlink to the paragraph
            paragraph._p.append(hyperlink)
            
            return text
        except Exception as e:
            # Fallback if hyperlink creation fails
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            run.font.underline = True
            return text
    
    def update_progress(self, value, text):
        """Update progress bar and label"""
        self.progress_bar.set(value)
        self.progress_label.configure(text=text)
        self.root.update_idletasks()
    
    def process_resume_thread(self):
        """Process resume in a separate thread to keep UI responsive"""
        # Get inputs
        job_url = self.url_entry.get().strip()
        job_desc_text = self.job_desc_text.get("1.0", tk.END).strip()
        resume_path = self.file_entry.get().strip()
        
        # Show progress indicator
        self.progress_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.update_progress(0.05, "Starting optimization process...")
        
        # Validate inputs
        if not resume_path or not os.path.exists(resume_path):
            messagebox.showerror("Error", "Please select a valid resume file")
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        # Get job description (either from URL or text input)
        self.update_progress(0.1, "Obtaining job description...")
        
        if job_desc_text:
            # Use directly entered text
            job_description = job_desc_text
        elif job_url:
            # Extract from URL
            job_description = self.extract_job_description(job_url)
            if not job_description:
                self.progress_frame.pack_forget()
                self.process_button.configure(state="normal")
                return
        else:
            messagebox.showerror("Error", "Please enter either a job URL or paste the job description")
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        self.job_description = job_description
        
        # Extract resume text
        self.update_progress(0.3, "Reading and analyzing resume file...")
        resume_text = ""
        
        try:
            if resume_path.lower().endswith('.pdf'):
                with open(resume_path, 'rb') as file:
                    resume_text, self.original_resume_content = self.extract_text_from_pdf(file)
            elif resume_path.lower().endswith('.docx'):
                with open(resume_path, 'rb') as file:
                    resume_text, self.original_resume_content = self.extract_text_from_docx(file)
            else:
                messagebox.showerror("Error", "Unsupported file format. Please use PDF or DOCX files.")
                self.progress_frame.pack_forget()
                self.process_button.configure(state="normal")
                return
        except Exception as e:
            messagebox.showerror("Error", f"Error reading resume file: {e}")
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        if not resume_text:
            messagebox.showerror("Error", "Could not extract text from the resume")
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        # Process with OpenAI
        self.update_progress(0.5, "Optimizing resume... (this may take a minute)")
        optimized_resume = self.optimize_resume(resume_text, job_description)
        
        if not optimized_resume:
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        self.optimized_resume_text = optimized_resume
        
        self.update_progress(0.8, "Generating analysis report...")
        analysis_report = self.generate_analysis_report(resume_text, job_description)
        
        if not analysis_report:
            self.progress_frame.pack_forget()
            self.process_button.configure(state="normal")
            return
        
        self.analysis_report = analysis_report
        
        # Update UI with results
        self.update_progress(1.0, "Completed!")
        self.root.after(0, self.update_results, optimized_resume, analysis_report)
    
    def update_results(self, optimized_resume, analysis_report):
        """Update UI with results (called from main thread)"""
        # Clear previous results
        self.analysis_text.delete("1.0", tk.END)
        self.resume_text.delete("1.0", tk.END)
        
        # Insert new results
        self.analysis_text.insert("1.0", analysis_report)
        self.resume_text.insert("1.0", optimized_resume)
        
        # Format text display
        self.apply_text_highlighting(self.analysis_text)
        
        # Show results and save sections
        self.results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.save_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Update progress and button
        self.progress_label.configure(text="Resume optimization complete!")
        self.process_button.configure(state="normal")
        
        # Switch to analysis tab
        self.tab_view.set("Analysis Report")
        
        # Scroll to top of results
        self.analysis_text.see("1.0")
        self.resume_text.see("1.0")
    
    def apply_text_highlighting(self, text_widget):
        """Apply formatting to the analysis text"""
        content = text_widget.get("1.0", tk.END)
        
        # Define tags
        text_widget.tag_configure("heading", font=("Arial", 14, "bold"))
        text_widget.tag_configure("subheading", font=("Arial", 12, "bold"), foreground="#1E88E5")
        text_widget.tag_configure("bullet", foreground="#26A69A")
        text_widget.tag_configure("highlight", background="#FFF9C4")
        
        # Clear text and reinsert with formatting
        text_widget.delete("1.0", tk.END)
        text_widget.insert(tk.END, content)
        
        # Apply formatting
        lines = content.split('\n')
        pos = "1.0"
        
        for line in lines:
            line_end = f"{pos}+{len(line)}c"
            
            # Headings (numbered items)
            if re.match(r'^\d+\.\s+\w+', line.strip()):
                text_widget.tag_add("heading", pos, line_end)
            
            # Subheadings (terms like "Strengths:", "Weaknesses:", etc.)
            if line.strip().endswith(':'):
                text_widget.tag_add("subheading", pos, line_end)
            
            # Bullets
            if line.strip().startswith('•') or line.strip().startswith('-'):
                text_widget.tag_add("bullet", pos, f"{pos}+1c")
            
            # Next line
            pos = f"{line_end}+1c"
    
    def process_resume(self):
        """Start resume processing"""
        # Disable process button while working
        self.process_button.configure(state="disabled")
        
        # Start processing in a separate thread
        threading.Thread(target=self.process_resume_thread, daemon=True).start()

def main():
    # Check if running as script or frozen executable
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        os.chdir(os.path.dirname(sys.executable))
    else:
        # Running as script
        os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Create root window
    root = ctk.CTk()
    app = ResumeOptimizerApp(root)
    root.mainloop()

if __name__ == "__main__":
    import sys
    main()